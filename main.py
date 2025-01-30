import streamlit as st
from PIL import Image
from docx import Document

# Configurações da Página
st.set_page_config(
    page_title="Robô Especializado",
    page_icon="🤖",
    layout="centered"
)

# Função para Gerar Termo de Referência
def gerar_tr(dados, tabela_precos):
    doc = Document()
    doc.add_heading('Termo de Referência', level=1)

    for secao, conteudo in dados.items():
        doc.add_heading(secao, level=2)
        doc.add_paragraph(conteudo if conteudo else "Não informado")

    # Adicionar Tabela de Preços
    doc.add_heading("13. Estimativa de Preço", level=2)
    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Descrição"
    hdr_cells[2].text = "Quantidade"
    hdr_cells[3].text = "Valor Unitário"
    hdr_cells[4].text = "Valor Total"
    
    for linha in tabela_precos:
        row_cells = tabela.add_row().cells
        row_cells[0].text = str(linha["Item"])
        row_cells[1].text = linha["Descrição"]
        row_cells[2].text = str(linha["Quantidade"])
        row_cells[3].text = f'R$ {linha["Valor Unitário"]:.2f}'
        row_cells[4].text = f'R$ {linha["Valor Total"]:.2f}'

    doc.save("Termo_de_Referencia.docx")
    return "Termo_de_Referencia.docx"

# Função da Página Principal
def main_page():
    st.markdown(
        """<h1 style='text-align: center; color: #4CAF50;'>
        🤖 Robô Especializado
        </h1>""",
        unsafe_allow_html=True
    )

    st.markdown(
        """<p style='text-align: center; font-size: 18px;'>
        Estou aqui para ajudar você a simplificar tarefas e otimizar processos. Escolha abaixo o sistema que deseja acessar:
        </p>""",
        unsafe_allow_html=True
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        st.image("https://via.placeholder.com/150", caption="Base de Conhecimento", use_container_width=True)
        if st.button("Base de Conhecimento"):
            st.session_state.page = "base_conhecimento"

    with col2:
        st.image("https://via.placeholder.com/150", caption="Chat GPT", use_container_width=True)
        if st.button("Chat GPT"):
            st.session_state.page = "chat_gpt"

    with col3:
        st.image("https://via.placeholder.com/150", caption="Termo de Referência", use_container_width=True)
        if st.button("Termo de Referência"):
            st.session_state.page = "termo_referencia"

# Página do Gerador de TR
def termo_referencia_page():
    st.title("Gerador de Termo de Referência")
    st.write("Preencha os campos abaixo para gerar o Termo de Referência baseado na Lei 14.133/2021.")

    with st.form(key="formulario_tr"):
        definicao_objeto = st.text_area("1. Definição do Objeto")
        justificativa = st.text_area("2. Justificativa")
        fundamentacao_legal = st.text_area("3. Fundamentação Legal")
        especificacoes_tecnicas = st.text_area("4. Especificações Técnicas")
        criterios_aceitacao = st.text_area("5. Critérios de Aceitação")
        resultados_esperados = st.text_area("6. Resultados Esperados")
        quantidade_prazo = st.text_area("7. Quantidade e Prazo")
        vigencia_contrato = st.text_area("8. Vigência do Contrato")
        criterios_selecao = st.text_area("9. Critérios de Seleção")
        fontes_recursos = st.text_area("10. Fontes de Recursos")
        plano_fiscalizacao = st.text_area("11. Plano de Fiscalização")
        anexos = st.text_area("12. Anexos")

        tabela_precos = st.experimental_data_editor(
            {"Item": [1], "Descrição": [""], "Quantidade": [1], "Valor Unitário": [0.0], "Valor Total": [0.0]},
            num_rows="dynamic",
        )

        submit_button = st.form_submit_button(label="Gerar Termo de Referência")

    if submit_button:
        dados = {
            "1. Definição do Objeto": definicao_objeto,
            "2. Justificativa": justificativa,
            "3. Fundamentação Legal": fundamentacao_legal,
            "4. Especificações Técnicas": especificacoes_tecnicas,
            "5. Critérios de Aceitação": criterios_aceitacao,
            "6. Resultados Esperados": resultados_esperados,
            "7. Quantidade e Prazo": quantidade_prazo,
            "8. Vigência do Contrato": vigencia_contrato,
            "9. Critérios de Seleção": criterios_selecao,
            "10. Fontes de Recursos": fontes_recursos,
            "11. Plano de Fiscalização": plano_fiscalizacao,
            "12. Anexos": anexos,
        }
        arquivo = gerar_tr(dados, tabela_precos)
        st.success("O Termo de Referência foi gerado com sucesso!")
        st.download_button("Baixar Termo de Referência", data=open(arquivo, "rb"), file_name=arquivo)

# Controle de Navegação
if "page" not in st.session_state:
    st.session_state.page = "main"

if st.session_state.page == "main":
    main_page()
elif st.session_state.page == "termo_referencia":
    termo_referencia_page()
