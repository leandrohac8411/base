import streamlit as st
from docx import Document

# Função para gerar o TR
def gerar_tr(dados, tabela_precos):
    doc = Document()
    doc.add_heading('Termo de Referência', level=1)

    for secao, conteudo in dados.items():
        doc.add_heading(secao, level=2)
        doc.add_paragraph(conteudo if conteudo else "Não informado")

    # Adicionar a tabela de preços
    doc.add_heading("13. Estimativa de Preço", level=2)
    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'

    # Cabeçalhos da tabela
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Descrição"
    hdr_cells[2].text = "Quantidade"
    hdr_cells[3].text = "Valor Unitário"
    hdr_cells[4].text = "Valor Total"

    # Preenchendo a tabela com os dados
    for linha in tabela_precos:
        row_cells = tabela.add_row().cells
        row_cells[0].text = str(linha["Item"])
        row_cells[1].text = linha["Descrição"]
        row_cells[2].text = str(linha["Quantidade"])
        row_cells[3].text = f'R$ {linha["Valor Unitário"]:.2f}'
        row_cells[4].text = f'R$ {linha["Valor Total"]:.2f}'

    doc.save("Termo_de_Referencia.docx")
    return "Termo_de_Referencia.docx"

# Configuração do Streamlit
st.title("Gerador de Termo de Referência")
st.write("Preencha os campos abaixo para gerar o Termo de Referência baseado na Lei 14.133/2021.")

# Formulário para entrada de dados
with st.form(key="formulario_tr"):
    definicao_objeto = st.text_area("1. Definição do Objeto", placeholder="Descreva o objeto a ser contratado...")
    justificativa = st.text_area("2. Justificativa", placeholder="Explique a necessidade e importância da contratação...")
    fundamentacao_legal = st.text_area("3. Fundamentação Legal", placeholder="Cite as bases legais aplicáveis...")
    especificacoes_tecnicas = st.text_area("4. Especificações Técnicas", placeholder="Detalhe as características técnicas do objeto...")
    criterios_aceitacao = st.text_area("5. Critérios de Aceitação", placeholder="Defina os parâmetros para aceitação...")
    resultados_esperados = st.text_area("6. Resultados Esperados", placeholder="Quais resultados a contratação deve gerar?")
    quantidade_prazo = st.text_area("7. Quantidade e Prazo", placeholder="Especifique as quantidades e prazos...")
    vigencia_contrato = st.text_area("8. Vigência do Contrato", placeholder="Informe a duração do contrato...")
    criterios_selecao = st.text_area("9. Critérios de Seleção", placeholder="Defina os critérios de escolha da proposta...")
    fontes_recursos = st.text_area("10. Fontes de Recursos", placeholder="Indique a origem dos recursos financeiros...")
    plano_fiscalizacao = st.text_area("11. Plano de Fiscalização", placeholder="Defina como será monitorada a execução do contrato...")
    anexos = st.text_area("12. Anexos", placeholder="Liste documentos complementares, se houver...")

    # Campo para tabela de estimativa de preços
    st.write("13. Estimativa de Preço")
    tabela_precos = []
    num_itens = st.number_input("Número de itens", min_value=1, step=1, value=1)
    for i in range(num_itens):
        st.write(f"Item {i + 1}")
        descricao = st.text_input(f"Descrição do item {i + 1}", key=f"desc_{i}")
        quantidade = st.number_input(f"Quantidade do item {i + 1}", min_value=1, step=1, value=1, key=f"quant_{i}")
        valor_unitario = st.number_input(f"Valor unitário do item {i + 1}", min_value=0.0, step=0.01, value=0.0, key=f"valor_{i}")
        valor_total = quantidade * valor_unitario
        tabela_precos.append({
            "Item": i + 1,
            "Descrição": descricao,
            "Quantidade": quantidade,
            "Valor Unitário": valor_unitario,
            "Valor Total": valor_total
        })
        st.write(f"Valor Total do item {i + 1}: R$ {valor_total:.2f}")

    # Botão para enviar o formulário
    submit_button = st.form_submit_button(label="Gerar Termo de Referência")

# Geração do documento
if submit_button:
    dados = {
        "1. Definição do Objeto": definicao_objeto,
        "2. Justificativa": justificativa,
        "3. Fundamentação Legal": fundamentacao_legal,
        "4. Especificações Técnicas": especificacoes_tecnicas,
        "5. Critérios de Aceitação": criterios_aceitacao,
        "6. Resultados Esperados": resultados_esperados,
        "7. Quantidade e Prazo": quantidade_prazo,
        "8. Vigência do Contrato": vigencia_contrato,
        "9. Critérios de Seleção": criterios_selecao,
        "10. Fontes de Recursos": fontes_recursos,
        "11. Plano de Fiscalização": plano_fiscalizacao,
        "12. Anexos": anexos,
    }
    arquivo = gerar_tr(dados, tabela_precos)
    st.success("O Termo de Referência foi gerado com sucesso!")
    st.download_button("Baixar Termo de Referência", data=open(arquivo, "rb"), file_name=arquivo)
    st.experimental_rerun()
