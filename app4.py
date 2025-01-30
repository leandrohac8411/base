import streamlit as st
from docx import Document
from io import BytesIO
import os
import re

# Função para detectar campos com marcadores no formato <escrever>
def detectar_campos(modelo_path):
    doc = Document(modelo_path)
    campos = []
    for p in doc.paragraphs:
        matches = re.finditer(r"<(.*?)>", p.text)
        for match in matches:
            campo = match.group()
            contexto = p.text.strip()
            if len(contexto) > 50:
                contexto = f"{contexto[:50]}..."  # Limitar o contexto para leitura
            campos.append({"campo": campo, "contexto": contexto})
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.finditer(r"<(.*?)>", cell.text)
                for match in matches:
                    campo = match.group()
                    contexto = cell.text.strip()
                    if len(contexto) > 50:
                        contexto = f"{contexto[:50]}..."  # Limitar o contexto
                    campos.append({"campo": campo, "contexto": contexto})
    # Remover duplicatas
    campos_unicos = []
    vistos = set()
    for campo in campos:
        if campo["campo"] not in vistos:
            campos_unicos.append(campo)
            vistos.add(campo["campo"])
    return campos_unicos

# Preencher o documento com valores fornecidos
def preencher_documento(modelo_path, campos_preenchidos):
    doc = Document(modelo_path)
    for p in doc.paragraphs:
        for campo, valor in campos_preenchidos.items():
            if campo in p.text:
                p.text = p.text.replace(campo, valor)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for campo, valor in campos_preenchidos.items():
                    if campo in cell.text:
                        cell.text = cell.text.replace(campo, valor)
    return doc

# Salvar o documento preenchido como um arquivo
def salvar_documento(doc, nome_arquivo="documento_preenchido.docx"):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Interface do Streamlit
st.set_page_config(page_title="Preenchimento Automático de Documentos 📄", page_icon="📄")
st.title("Preenchimento Automático de Documentos 📄")

# Envio do modelo de documento
st.sidebar.header("Envie o modelo de documento")
modelo_file = st.sidebar.file_uploader("Envie o modelo de documento (DOCX)", type=["docx"])

if modelo_file:
    modelo_path = os.path.join("temp_model.docx")
    with open(modelo_path, "wb") as f:
        f.write(modelo_file.getvalue())

    st.sidebar.success("Modelo carregado com sucesso!")

    # Detectar campos no modelo
    campos_detectados = detectar_campos(modelo_path)

    if campos_detectados:
        st.header("Preencha os Campos")
        campos_preenchidos = {}
        for idx, item in enumerate(campos_detectados):
            campo = item["campo"]
            contexto = item["contexto"]
            campos_preenchidos[campo] = st.text_input(f"Campo {idx + 1} - {contexto}", value="", key=f"campo_{idx}")

        if st.button("Gerar Documento"):
            if all(value.strip() for value in campos_preenchidos.values()):
                # Preencher o documento com os valores inseridos
                doc_preenchido = preencher_documento(modelo_path, campos_preenchidos)
                buffer = salvar_documento(doc_preenchido)

                st.download_button(
                    label="Baixar Documento Preenchido",
                    data=buffer,
                    file_name="documento_preenchido.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.error("Por favor, preencha todos os campos antes de gerar o documento.")
    else:
        st.warning("Nenhum campo identificado automaticamente. Ajuste o modelo para incluir marcadores como '<escrever>'.")
else:
    st.info("Envie um modelo de documento no formato DOCX para começar.")
