import streamlit as st
from docx import Document
from io import BytesIO
import os
import re

# Função para detectar campos com contexto no documento
def detectar_campos_com_contexto(modelo_path):
    doc = Document(modelo_path)
    campos = []
    for p in doc.paragraphs:
        matches = re.finditer(r"(_{3,}|<.*?>|\[.*?\])", p.text)
        for match in matches:
            campo = match.group()
            contexto = p.text.strip()  # Extrair o texto completo do parágrafo
            if len(contexto) > 50:  # Limitar o tamanho do contexto para exibição
                contexto = f"{contexto[:50]}..."
            campos.append({"campo": campo, "contexto": contexto})
    return campos

# Preencher o documento com valores fornecidos
def preencher_documento(modelo_path, campos_preenchidos):
    doc = Document(modelo_path)
    for p in doc.paragraphs:
        for campo, valor in campos_preenchidos.items():
            p.text = p.text.replace(campo, valor)
    return doc

# Salvar o documento preenchido como um arquivo
def salvar_documento(doc, nome_arquivo="documento_preenchido.docx"):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Interface do Streamlit
st.set_page_config(page_title="Preenchimento Automático de Documentos 📄", page_icon="📄")
st.title("Preenchimento Automático de Documentos 📄")

# Envio do modelo de documento
st.sidebar.header("Envie o modelo de documento")
modelo_file = st.sidebar.file_uploader("Envie o modelo de documento (DOCX)", type=["docx"])

if modelo_file:
    modelo_path = os.path.join("temp_model.docx")
    with open(modelo_path, "wb") as f:
        f.write(modelo_file.getvalue())

    st.sidebar.success("Modelo carregado com sucesso!")

    # Detectar campos no modelo com contexto
    campos_detectados = detectar_campos_com_contexto(modelo_path)

    if campos_detectados:
        st.header("Preencha os Campos")
        campos_preenchidos = {}
        for item in campos_detectados:
            campo = item["campo"]
            contexto = item["contexto"]
            campos_preenchidos[campo] = st.text_input(f"Preencha o campo ({contexto}):", value="")

        if st.button("Gerar Documento"):
            doc_preenchido = preencher_documento(modelo_path, campos_preenchidos)
            buffer = salvar_documento(doc_preenchido)

            st.download_button(
                label="Baixar Documento Preenchido",
                data=buffer,
                file_name="documento_preenchido.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    else:
        st.warning("Nenhum campo identificado automaticamente. Ajuste o modelo para incluir marcadores como '___', '<campo>' ou '[campo]'.")
else:
    st.info("Envie um modelo de documento no formato DOCX para começar.")
